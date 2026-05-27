import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\elvis\OneDrive\Documents\Final year project (MINETRACE)\MineTrace2.0 (1).docx')

# --- FIX 1-3: Paragraph fixes ---
fixes = [
    (335,
     'Notifications must be configurable and delivered via in-app alerts and email (where enabled).',
     'Notifications are stored in the system and displayed to relevant users via in-app alerts upon login or dashboard access.'),
    (323,
     'Automated alerts will notify relevant users when anomalies occur or when traceability steps are missing.',
     'The system records alerts for relevant users when anomalies occur or when traceability steps are missing.'),
    (553,
     'Users, Organizations, Mines, Mineral Batches, Movement Events, Risk Assessments, Verification Logs, and Audit Logs',
     'Users, Organizations, Mines, Mineral Batches, Movement Events, Verification Logs, and Audit Logs. Fraud risk data (flags, anomaly score, risk level) is stored directly within the Mineral Batch entity rather than as a separate table'),
]
for idx, old, new in fixes:
    para = doc.paragraphs[idx]
    full = ''.join(r.text for r in para.runs)
    if old in full:
        new_full = full.replace(old, new)
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ''
        print(f'Fixed para {idx}')
    elif old in para.text:
        # text may span runs differently
        para.runs[0].text = para.text.replace(old, new)
        for r in para.runs[1:]:
            r.text = ''
        print(f'Fixed para {idx} via para.text')
    else:
        print(f'Para {idx} NOT matched - text: {para.text[:80]}')

# --- FIX 4: Table 3 UC-03 event types ---
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            full_cell = ''.join(p.text for p in cell.paragraphs)
            if 'Dispatch/Receive/Store/Process/Sell' in full_cell:
                for para in cell.paragraphs:
                    full_p = ''.join(r.text for r in para.runs)
                    if 'Dispatch/Receive/Store/Process/Sell' in full_p:
                        new_p = full_p.replace('Dispatch/Receive/Store/Process/Sell', 'Dispatch/Receive/Storage/Transfer/Sale')
                        if para.runs:
                            para.runs[0].text = new_p
                            for r in para.runs[1:]:
                                r.text = ''
                        print('Fixed UC-03 event types')

# --- Helper: add a row to a table ---
def add_row(table, values):
    last_tr = list(table.rows)[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    tcs = new_tr.findall(qn('w:tc'))
    for i, val in enumerate(values):
        if i >= len(tcs):
            break
        tc = tcs[i]
        for p in tc.findall(qn('w:p')):
            tc.remove(p)
        p = OxmlElement('w:p')
        r_el = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = val
        r_el.append(t)
        p.append(r_el)
        tc.append(p)
    table._tbl.append(new_tr)

# --- FIX 5: Table 9 (Batch) - remove qr_code_path, add new rows ---
batch_table = doc.tables[9]
for row in list(batch_table.rows):
    if row.cells[0].text.strip() == 'qr_code_path':
        batch_table._tbl.remove(row._tr)
        print('Removed qr_code_path row from Batch table')
        break

add_row(batch_table, ['risk_level', 'Varchar(20)', '', 'Risk classification (Low/Medium/High/Unknown)'])
add_row(batch_table, ['anomaly_score', 'Decimal', '', 'Computed fraud anomaly score'])
add_row(batch_table, ['override_note', 'Text', '', 'Admin note when risk level is manually overridden'])
add_row(batch_table, ['inspector_approved', 'Boolean', '', 'Whether the inspector approved this batch'])
add_row(batch_table, ['inspector_note', 'Text', '', 'Inspector compliance note'])
print('Added 5 new rows to Batch table')

# --- FIX 6: Table 10 (Movement) - fix recorded_weight -> weight, recorded_by FK -> Varchar ---
movement_table = doc.tables[10]
for row in list(movement_table.rows):
    c0 = row.cells[0].text.strip()
    if c0 == 'recorded_weight':
        for para in row.cells[0].paragraphs:
            full_p = ''.join(r.text for r in para.runs)
            if para.runs:
                para.runs[0].text = full_p.replace('recorded_weight', 'weight')
                for r in para.runs[1:]:
                    r.text = ''
        print('Fixed recorded_weight -> weight in cell 0')
    elif c0 == 'recorded_by':
        # Change data type from Integer to Varchar(100)
        for para in row.cells[1].paragraphs:
            if para.runs:
                para.runs[0].text = 'Varchar(100)'
                for r in para.runs[1:]:
                    r.text = ''
        # Remove FK key type
        for para in row.cells[2].paragraphs:
            if para.runs:
                para.runs[0].text = ''
                for r in para.runs[1:]:
                    r.text = ''
        # Update description
        for para in row.cells[3].paragraphs:
            if para.runs:
                para.runs[0].text = 'Name of the user who recorded the movement'
                for r in para.runs[1:]:
                    r.text = ''
        print('Fixed recorded_by type and key')

# --- FIX 7: Table 11 (Fraud Analysis) - update to reflect embedded-in-Batch design ---
fraud_table = doc.tables[11]
fraud_rows = list(fraud_table.rows)

# Update header description column
header_cell = fraud_rows[0].cells[3]
for para in header_cell.paragraphs:
    if para.runs:
        para.runs[0].text = 'Description (all fields embedded in batches table)'
        for r in para.runs[1:]:
            r.text = ''

# Remove analysis_id and batch_id rows (replace with embedded-note row)
rows_to_remove = []
for row in fraud_rows[1:]:
    c0 = row.cells[0].text.strip()
    if c0 in ('analysis_id', 'batch_id'):
        rows_to_remove.append(row)
for row in rows_to_remove:
    fraud_table._tbl.remove(row._tr)
print(f'Removed {len(rows_to_remove)} rows from fraud table')

# Add clarifying rows at the end
add_row(fraud_table, ['Note', '', '', 'All fraud analysis fields are embedded directly in the Mineral Batch (batches table), not stored in a separate table.'])
add_row(fraud_table, ['batch_id', 'Integer', 'FK', 'Links fraud data to the batch being analyzed (same row in batches table)'])
add_row(fraud_table, ['handover_flag', 'Boolean', '', 'Indicates a missing handover step in the movement chain'])
add_row(fraud_table, ['weight_loss_flag', 'Boolean', '', 'Indicates significant weight drop between consecutive movements'])
add_row(fraud_table, ['future_extraction_flag', 'Boolean', '', 'Indicates the extraction date is set in the future (suspicious)'])
print('Updated Fraud Analysis table')

doc.save(r'C:\Users\elvis\OneDrive\Documents\Final year project (MINETRACE)\MineTrace2.0 (1).docx')
print('All fixes saved successfully.')
